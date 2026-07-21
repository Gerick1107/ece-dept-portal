"""Custom, template-driven publication exports.

The user uploads a template (CSV / XLSX / DOCX-with-a-table) whose header row
lists the columns they want, in the order they want. We map each requested
column to a known publication field, compile the matching data, and return a
file in the same format with exactly those columns.

Column matching is offline-first:
  1. exact / alias match (normalised),
  2. fuzzy match via difflib (surfaced for user confirmation),
  3. local LLM as a last resort for columns nothing else matched.

Columns that cannot be matched to any available field are reported back so the
caller can tell the user which data is missing.
"""
from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field as dc_field
from datetime import date
from difflib import SequenceMatcher

import pandas as pd
from sqlalchemy.orm import Session

from app.publications.exports.export_service import _faculty_names_map, _query_publications
from app.publications.models import Publication
from app.utils.docx_export import records_to_table_docx_bytes

# Canonical field -> human aliases the user might type in a template header.
CANONICAL_FIELDS: dict[str, list[str]] = {
    "title": ["title", "paper title", "publication title", "name", "paper name"],
    "authors": ["authors", "author", "author names", "author list"],
    "journal": ["journal", "journal name"],
    "conference": ["conference", "conference name", "proceedings"],
    "book": ["book", "book title"],
    "is_manual_book": ["manual book", "books tab", "is manual book", "assigned book"],
    "venue": ["venue", "published in", "source", "journal/conference", "journal or conference"],
    "volume": ["volume", "vol"],
    "issue": ["issue", "number", "issue no", "issue number"],
    "pages": ["pages", "page", "page numbers", "pp", "page range"],
    "publication_date": ["date", "publication date", "published date", "exact date", "date of publication"],
    "publication_year": ["year", "publication year", "pub year", "year of publication"],
    "publisher": ["publisher", "publishing house"],
    "citation_count": ["citations", "cited by", "citation count", "times cited", "no of citations"],
    "faculty_names": ["faculty", "faculty names", "faculty name", "ece faculty", "faculty authors"],
    "inventors": ["inventors", "inventor", "inventor names"],
    "patent_number": ["patent number", "patent no", "patent no."],
    "patent_office": ["patent office", "office"],
    "application_number": ["application number", "application no", "app number"],
    "link": ["link", "url", "web link", "paper link", "doi link"],
}

_FUZZY_THRESHOLD = 0.72


def _normalize(text: str) -> str:
    return "".join(ch for ch in str(text or "").strip().lower() if ch.isalnum() or ch == " ").strip()


# Precompute alias -> field lookup.
_ALIAS_TO_FIELD: dict[str, str] = {}
for _field, _aliases in CANONICAL_FIELDS.items():
    _ALIAS_TO_FIELD[_normalize(_field)] = _field
    for _alias in _aliases:
        _ALIAS_TO_FIELD[_normalize(_alias)] = _field


@dataclass
class MappingResult:
    # header -> field for confident (exact/alias) matches
    matched: dict[str, str] = dc_field(default_factory=dict)
    # header -> {field, score} for fuzzy matches needing confirmation
    suggestions: dict[str, dict] = dc_field(default_factory=dict)
    # headers that matched nothing
    unknown: list[str] = dc_field(default_factory=list)


class TemplateError(Exception):
    """Raised when a template cannot be parsed or a column cannot be satisfied."""


# --- Header extraction ------------------------------------------------------

def _detect_format(filename: str) -> str:
    lower = (filename or "").lower()
    if lower.endswith(".csv"):
        return "csv"
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        return "xlsx"
    if lower.endswith(".docx"):
        return "docx"
    raise TemplateError("Unsupported template type. Use a .csv, .xlsx, or .docx file.")


def parse_fields_from_text(text: str) -> list[str]:
    """Extract column names from free-form text (lists, tables, comma-separated, etc.)."""
    import re

    raw = (text or "").strip()
    if not raw:
        return []

    headers: list[str] = []
    seen_norm: set[str] = set()

    def _add(candidate: str) -> None:
        cleaned = re.sub(r"^[\d]+[\.\)\]:]\s*", "", candidate.strip())
        cleaned = re.sub(r"^[-•*]\s*", "", cleaned).strip()
        cleaned = cleaned.strip("\"'`")
        if not cleaned or len(cleaned) > 120:
            return
        norm = _normalize(cleaned)
        if norm and norm not in seen_norm:
            seen_norm.add(norm)
            headers.append(cleaned)

    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        if re.search(r"[,;\t|]", line):
            for part in re.split(r"[,;\t|]+", line):
                _add(part)
        else:
            _add(line)

    if not headers and raw:
        if re.search(r"[,;\t|]", raw):
            for part in re.split(r"[,;\t|]+", raw):
                _add(part)
        else:
            _add(raw)
    return headers


def extract_headers(filename: str, content: bytes) -> tuple[list[str], str]:
    fmt = _detect_format(filename)
    if fmt == "csv":
        text = content.decode("utf-8-sig", errors="replace")
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            headers = [c.strip() for c in row if c.strip()]
            if headers:
                return headers, fmt
        raise TemplateError("The CSV template has no header row.")
    if fmt == "xlsx":
        df = pd.read_excel(io.BytesIO(content), nrows=0)
        headers = [str(c).strip() for c in df.columns if str(c).strip() and not str(c).startswith("Unnamed")]
        if not headers:
            raise TemplateError("The Excel template's first row has no column names.")
        return headers, fmt
    # docx: read the first table's header row
    from docx import Document

    document = Document(io.BytesIO(content))
    if not document.tables:
        raise TemplateError("The Word template must contain a table whose first row lists the columns.")
    header_cells = document.tables[0].rows[0].cells
    headers = [c.text.strip() for c in header_cells if c.text.strip()]
    if not headers:
        raise TemplateError("The Word template's table has no header row.")
    return headers, fmt


# --- Column matching --------------------------------------------------------

def _fuzzy_field(header_norm: str) -> tuple[str | None, float]:
    best_field: str | None = None
    best_score = 0.0
    for alias_norm, field_name in _ALIAS_TO_FIELD.items():
        score = SequenceMatcher(None, header_norm, alias_norm).ratio()
        if score > best_score:
            best_score, best_field = score, field_name
    return best_field, best_score


def custom_field_aliases(db: Session | None) -> dict[str, str]:
    """Map normalized aliases -> ``custom:<key>`` for admin-defined columns."""
    if db is None:
        return {}
    from app.publications.services.custom_columns_service import list_columns

    aliases: dict[str, str] = {}
    for col in list_columns(db, enabled_only=True):
        target = f"custom:{col.key}"
        aliases[_normalize(col.label)] = target
        aliases[_normalize(col.key)] = target
    return aliases


def match_headers(headers: list[str], db: Session | None = None) -> MappingResult:
    result = MappingResult()
    alias_map = {**_ALIAS_TO_FIELD, **custom_field_aliases(db)}
    for header in headers:
        norm = _normalize(header)
        if norm in alias_map:
            result.matched[header] = alias_map[norm]
            continue
        field_name, score = _fuzzy_field(norm)
        if field_name and score >= _FUZZY_THRESHOLD:
            result.suggestions[header] = {"field": field_name, "score": round(score, 2)}
        else:
            result.unknown.append(header)
    return result


async def llm_guess_fields(unknown: list[str]) -> dict[str, str]:
    """Last-resort mapping of unknown headers to canonical fields via the local LLM."""
    if not unknown:
        return {}
    from app.llm.services.llm_dispatch import generate_text

    fields = ", ".join(CANONICAL_FIELDS.keys())
    prompt = (
        "Map each requested spreadsheet column to the single best matching field from this "
        f"fixed list (or 'none' if nothing fits): {fields}.\n"
        "Return one line per column as 'column => field'. Columns:\n"
        + "\n".join(f"- {h}" for h in unknown)
    )
    try:
        text = await generate_text(prompt, provider="local", temperature=0.0, max_tokens=400)
    except Exception:
        return {}
    guesses: dict[str, str] = {}
    for line in text.splitlines():
        if "=>" not in line:
            continue
        left, right = line.split("=>", 1)
        header = left.strip().lstrip("-").strip()
        field_name = _normalize(right)
        for original in unknown:
            if _normalize(original) == _normalize(header) and field_name in CANONICAL_FIELDS:
                guesses[original] = field_name
    return guesses


# --- Compilation ------------------------------------------------------------

def _field_value(row: Publication, field_name: str, faculty_names: str) -> str:
    if field_name.startswith("custom:"):
        from app.publications.services.custom_columns_service import get_custom_fields

        return get_custom_fields(row).get(field_name.split(":", 1)[1], "")
    if field_name == "venue":
        return row.journal or row.conference or row.book or ""
    if field_name == "faculty_names":
        return faculty_names
    if field_name == "link":
        return row.link or row.scholar_url or ""
    value = getattr(row, field_name, None)
    return "" if value is None else str(value)


def compile_export(
    db: Session,
    *,
    headers: list[str],
    mapping: dict[str, str],
    fmt: str,
    faculty_ids: list[int] | None = None,
    publication_year: int | None = None,
    year_start: int | None = None,
    year_end: int | None = None,
    date_start: date | None = None,
    date_end: date | None = None,
    export_type: str = "both",
    title: str = "Publications Export",
) -> bytes:
    """Build a file in ``fmt`` containing exactly ``headers`` (in order), using ``mapping``."""
    missing = [h for h in headers if h not in mapping]
    if missing:
        raise TemplateError(
            "These columns could not be matched to any available data: " + ", ".join(missing)
        )

    rows = _query_publications(
        db,
        faculty_ids=faculty_ids,
        publication_year=publication_year,
        year_start=year_start,
        year_end=year_end,
        date_start=date_start,
        date_end=date_end,
        export_type=export_type,
    )
    names_map = _faculty_names_map(db, [r.id for r in rows])

    records: list[dict] = []
    for row in rows:
        faculty_names = names_map.get(row.id, "")
        records.append({h: _field_value(row, mapping[h], faculty_names) for h in headers})

    if fmt == "csv":
        return pd.DataFrame(records, columns=headers).to_csv(index=False).encode("utf-8")
    if fmt == "xlsx":
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            pd.DataFrame(records, columns=headers).to_excel(writer, index=False, sheet_name="Publications")
        return output.getvalue()
    if fmt == "docx":
        return records_to_table_docx_bytes(title, headers, records)
    if fmt == "pdf":
        from app.utils.pdf_tables import records_to_list_pdf_bytes

        pdf_records = [{h: rec.get(h, "") for h in headers} for rec in records]
        return records_to_list_pdf_bytes(title, pdf_records)
    raise TemplateError(f"Unsupported output format: {fmt}")


def compile_fields_export(
    db: Session,
    *,
    headers: list[str],
    mapping: dict[str, str],
    fmt: str,
    faculty_ids: list[int] | None = None,
    publication_year: int | None = None,
    year_start: int | None = None,
    year_end: int | None = None,
    date_start: date | None = None,
    date_end: date | None = None,
    export_type: str = "both",
    title: str = "Publications Export",
) -> bytes:
    """Build export in the chosen format from a user-defined column list (no template file)."""
    if fmt not in ("csv", "xlsx", "pdf", "docx"):
        raise TemplateError("Unsupported format. Choose csv, xlsx, pdf, or docx.")
    if not headers:
        raise TemplateError("Add at least one column name.")
    return compile_export(
        db,
        headers=headers,
        mapping=mapping,
        fmt=fmt,
        faculty_ids=faculty_ids,
        publication_year=publication_year,
        year_start=year_start,
        year_end=year_end,
        date_start=date_start,
        date_end=date_end,
        export_type=export_type,
        title=title,
    )
