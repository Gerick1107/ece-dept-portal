"""Word (.docx) export helpers, mirroring ``pdf_tables`` for tabular data.

Uses python-docx. Provides a table renderer (fixed column order) and a
labeled-block renderer for records with long free-text fields.
"""
from __future__ import annotations

from io import BytesIO


def records_to_table_docx_bytes(
    title: str,
    columns: list[str],
    records: list[dict],
    *,
    header_labels: dict[str, str] | None = None,
) -> bytes:
    """Render records as a single Word table with ``columns`` as the header row/order."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading(title, level=1)

    labels = header_labels or {}
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(columns):
        run = header_cells[idx].paragraphs[0].add_run(labels.get(col, col.replace("_", " ").title()))
        run.bold = True
        run.font.size = Pt(10)

    for record in records:
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = record.get(col)
            cells[idx].text = "" if value is None else str(value)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def records_to_blocks_docx_bytes(
    title: str,
    records: list[dict],
    *,
    entry_title_key: str | None = None,
) -> bytes:
    """Render each record as a labeled block (good for long text fields)."""
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)

    for i, record in enumerate(records, start=1):
        heading = (
            f"{i}. {record[entry_title_key]}"
            if entry_title_key and record.get(entry_title_key)
            else f"Entry {i}"
        )
        document.add_heading(heading, level=3)
        for label, value in record.items():
            if entry_title_key and label == entry_title_key:
                continue
            para = document.add_paragraph()
            run = para.add_run(f"{label.replace('_', ' ').title()}: ")
            run.bold = True
            para.add_run("" if value is None else str(value))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
