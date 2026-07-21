"""Analyze uploaded question papers with the local LLM and build component marks templates."""

from __future__ import annotations

import io
import json
import re
from collections import defaultdict
from dataclasses import dataclass

from app.copo.services.marks_template_builder import build_analyzed_component_workbook

_MAX_TEXT_CHARS = 12000
_CO_RE = re.compile(r"^CO\s*\d+$", re.IGNORECASE)
# Q1a, Q1(b), Q1-ii, 1a, Q7c → parent key Q1 / 1 / Q7
_LABEL_PARENT_RE = re.compile(
    r"^(?P<parent>Q?\d+)\s*(?:[\.\-_]?\s*(?:\([a-z0-9ivx]+\)|[a-z]|[ivx]+|\d+))?$",
    re.IGNORECASE,
)
_PART_SUFFIX_RE = re.compile(
    r"^(?P<parent>Q?\d+)\s*(?:[\.\-_]?\s*(?:\([a-z0-9ivx]+\)|[a-z]|[ivx]+|\d+))$",
    re.IGNORECASE,
)
# Strip phantom trailing digits: Q7c1 → Q7c, Q4ii2 → Q4ii
_PHANTOM_TRAILING_DIGIT_RE = re.compile(r"^(?P<base>(?:Q)?\d+[a-z]+)(?P<extra>\d+)$", re.IGNORECASE)
# [Q3] [CO2, CO4] or Q3 (CO1, CO2) or Question 1: CO1, CO2
_PAPER_Q_CO_RE = re.compile(
    r"(?:\[?\s*Q(?:uestion)?\s*(?P<num>\d+)\s*\]?)"
    r"\s*[\[(:]?\s*"
    r"(?P<cos>(?:CO\s*\d+\s*[;,/&]?[\s]*){1,8})"
    r"\s*[\])]?",
    re.IGNORECASE,
)


@dataclass
class AnalyzedQuestion:
    label: str
    co_labels: list[str]
    max_marks: float
    is_bonus: bool = False

    @property
    def co_label(self) -> str:
        return _format_co_cell(self.co_labels)


def extract_document_text(filename: str, content: bytes) -> str:
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(parts).strip()
    if lower.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts).strip()
    if lower.endswith(".txt"):
        return content.decode("utf-8", errors="replace").strip()
    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def _parse_llm_json(text: str) -> dict:
    cleaned = (text or "").strip()
    if not cleaned:
        raise ValueError("LLM returned empty analysis. Try again.")
    if "```" in cleaned:
        match = re.search(r"```(?:json)?\s*([\s\S]*?)```", cleaned)
        if match:
            cleaned = match.group(1).strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        cleaned = cleaned[start : end + 1]
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise ValueError(
            "Could not parse LLM analysis JSON (response may have been truncated). "
            "Try again, or shorten the paper text."
        ) from exc


def _normalize_co(value: str) -> str:
    value = (value or "").strip().upper().replace(" ", "")
    if not value:
        return ""
    if value.startswith("CO") and value[2:].isdigit():
        return f"CO{int(value[2:])}"
    return value if _CO_RE.match(value) else ""


def _normalize_co_list(raw) -> list[str]:
    if raw is None:
        return []
    if isinstance(raw, list):
        labels: list[str] = []
        for item in raw:
            co = _normalize_co(str(item))
            if co and co not in labels:
                labels.append(co)
        return labels
    text = str(raw).strip()
    if not text:
        return []
    labels = []
    for part in re.split(r"[,;/&]+", text):
        co = _normalize_co(part)
        if co and co not in labels:
            labels.append(co)
    return labels


def _format_co_cell(labels: list[str]) -> str:
    return ", ".join(labels)


def extract_paper_co_map(text: str) -> dict[str, list[str]]:
    """
    Deterministically map question numbers → CO labels from the paper text.

    Handles common ECE formats like ``[Q3] [CO2, CO4]:`` so COs are not lost
    when the LLM omits or mis-quotes evidence.
    """
    mapping: dict[str, list[str]] = {}
    for match in _PAPER_Q_CO_RE.finditer(text or ""):
        num = match.group("num")
        cos = _normalize_co_list(match.group("cos"))
        if not num or not cos:
            continue
        key = f"Q{int(num)}"
        existing = mapping.setdefault(key, [])
        for co in cos:
            if co not in existing:
                existing.append(co)
    return mapping


def _cos_for_label(label: str, paper_co_map: dict[str, list[str]]) -> list[str]:
    parent = _parent_key(label)
    if not parent:
        return []
    # Normalize 3 → Q3
    if parent.isdigit():
        parent = f"Q{int(parent)}"
    elif parent.upper().startswith("Q") and parent[1:].isdigit():
        parent = f"Q{int(parent[1:])}"
    return list(paper_co_map.get(parent.upper(), []) or paper_co_map.get(parent, []))


async def _generate_analysis_text(prompt: str) -> str:
    from app.llm.services.llm_dispatch import generate_text

    # Papers with many sub-parts need headroom to avoid truncated JSON.
    return await generate_text(prompt, provider="local", temperature=0.0, max_tokens=4096)


def _validated_co_labels(item: dict, source_text: str) -> list[str]:
    """Accept LLM CO labels only when backed by an exact quote from the paper."""
    labels = _normalize_co_list(item.get("co_labels"))
    if not labels:
        labels = _normalize_co_list(item.get("co_label"))
    evidence = str(item.get("co_evidence") or "").strip()
    if not labels or not evidence:
        return []

    normalized_source = " ".join(source_text.split()).casefold()
    normalized_evidence = " ".join(evidence.split()).casefold()
    if normalized_evidence not in normalized_source:
        return []

    evidence_labels = _normalize_co_list(
        re.findall(r"\bCO\s*\d+\b", evidence, flags=re.IGNORECASE)
    )
    return [label for label in labels if label in evidence_labels]


def _safe_float(value, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _equal_split(total: float, n: int) -> list[float]:
    """Split total into n parts that sum exactly to total (2-decimal marks)."""
    if n <= 0:
        return []
    if total <= 0:
        return [0.0] * n
    each = round(total / n, 2)
    parts = [each] * (n - 1)
    parts.append(round(total - each * (n - 1), 2))
    return parts


def _part_marks_need_equal_split(parent_marks: float, part_marks: list[float]) -> bool:
    """True when part-wise marks are missing or each part copied the parent total."""
    if not part_marks:
        return False
    if parent_marks <= 0:
        return False
    if all(m <= 0 for m in part_marks):
        return True
    # Classic LLM bug: every sub-part gets the parent's total (e.g. all 10).
    if len(part_marks) > 1 and all(abs(m - parent_marks) < 1e-6 for m in part_marks):
        return True
    return False


def _parent_key(label: str) -> str | None:
    text = (label or "").strip()
    match = _LABEL_PARENT_RE.match(text)
    if not match:
        return None
    return match.group("parent").upper()


def _is_subpart_label(label: str) -> bool:
    return bool(_PART_SUFFIX_RE.match((label or "").strip()))


def normalize_question_label(label: str) -> str:
    """Fix common LLM phantoms like Q7c1 → Q7c."""
    text = (label or "").strip()
    if not text:
        return text
    match = _PHANTOM_TRAILING_DIGIT_RE.match(text)
    if match:
        return match.group("base")
    return text


def apply_paper_cos(
    questions: list[AnalyzedQuestion],
    paper_co_map: dict[str, list[str]],
) -> list[AnalyzedQuestion]:
    """Prefer COs parsed from the paper; fall back to whatever the LLM already set."""
    if not paper_co_map:
        return questions
    updated: list[AnalyzedQuestion] = []
    for q in questions:
        paper_cos = _cos_for_label(q.label, paper_co_map)
        cos = paper_cos or q.co_labels
        updated.append(AnalyzedQuestion(q.label, cos, q.max_marks, is_bonus=q.is_bonus))
    return updated


def dedupe_questions(questions: list[AnalyzedQuestion]) -> list[AnalyzedQuestion]:
    """Merge duplicate labels after normalization (keep first non-zero marks / COs)."""
    by_label: dict[str, AnalyzedQuestion] = {}
    order: list[str] = []
    for q in questions:
        label = normalize_question_label(q.label)
        key = label.upper()
        if key not in by_label:
            by_label[key] = AnalyzedQuestion(label, list(q.co_labels), q.max_marks, is_bonus=q.is_bonus)
            order.append(key)
            continue
        existing = by_label[key]
        cos = existing.co_labels or q.co_labels
        marks = existing.max_marks if existing.max_marks > 0 else q.max_marks
        by_label[key] = AnalyzedQuestion(
            existing.label,
            cos,
            marks,
            is_bonus=existing.is_bonus or q.is_bonus,
        )
    return [by_label[k] for k in order]


def redistribute_sibling_marks(
    questions: list[AnalyzedQuestion],
    *,
    paper_total_marks: float,
) -> list[AnalyzedQuestion]:
    """
    When the LLM assigns the parent question total to every sub-part (e.g. each of
    Q1a/Q1b marked 10 when the question is worth 10), divide that total equally.

    Only runs when the flat sum clearly overshoots paper_total_marks, so already-
    correct equal splits (5+5) are left alone.
    """
    if len(questions) < 2:
        return questions

    non_bonus = [q for q in questions if not q.is_bonus]
    non_bonus_sum = sum(q.max_marks for q in non_bonus)
    if paper_total_marks <= 0 or non_bonus_sum <= paper_total_marks * 1.15:
        return questions

    groups: dict[str, list[int]] = defaultdict(list)
    for idx, q in enumerate(questions):
        if q.is_bonus or not _is_subpart_label(q.label):
            continue
        key = _parent_key(q.label)
        if key:
            groups[key].append(idx)

    updated = list(questions)
    for indices in groups.values():
        if len(indices) < 2:
            continue
        marks = [updated[i].max_marks for i in indices]
        if not marks or min(marks) != max(marks) or marks[0] <= 0:
            continue
        splits = _equal_split(marks[0], len(indices))
        for i, split in zip(indices, splits):
            q = updated[i]
            updated[i] = AnalyzedQuestion(q.label, q.co_labels, split, is_bonus=q.is_bonus)
    return updated


def _flatten_llm_questions(data: dict, source_text: str) -> list[AnalyzedQuestion]:
    questions: list[AnalyzedQuestion] = []
    for item in data.get("questions") or []:
        label = normalize_question_label(str(item.get("label") or "").strip() or f"Q{len(questions) + 1}")
        parent_cos = _validated_co_labels(item, source_text)
        parent_marks = _safe_float(item.get("max_marks"))
        parent_bonus = bool(item.get("is_bonus"))
        raw_parts = item.get("parts") or []

        if raw_parts:
            part_specs: list[tuple[str, list[str], float, bool]] = []
            for part in raw_parts:
                plabel = normalize_question_label(str(part.get("label") or "").strip())
                if not plabel:
                    continue
                pcos = _validated_co_labels(part, source_text) or list(parent_cos)
                pmarks = _safe_float(part.get("max_marks"))
                part_specs.append((plabel, pcos, pmarks, bool(part.get("is_bonus")) or parent_bonus))

            if not part_specs:
                questions.append(AnalyzedQuestion(label, parent_cos, parent_marks, is_bonus=parent_bonus))
                continue

            part_mark_values = [p[2] for p in part_specs]
            if _part_marks_need_equal_split(parent_marks, part_mark_values):
                splits = _equal_split(parent_marks, len(part_specs))
            elif all(m <= 0 for m in part_mark_values) and parent_marks > 0:
                splits = _equal_split(parent_marks, len(part_specs))
            else:
                explicit = [m for m in part_mark_values if m > 0]
                if parent_marks > 0 and len(explicit) < len(part_specs):
                    used = sum(explicit)
                    remaining = max(parent_marks - used, 0.0)
                    zero_idxs = [i for i, m in enumerate(part_mark_values) if m <= 0]
                    fills = _equal_split(remaining, len(zero_idxs)) if zero_idxs else []
                    splits = list(part_mark_values)
                    for i, fill in zip(zero_idxs, fills):
                        splits[i] = fill
                else:
                    splits = part_mark_values

            for (plabel, pcos, _, pbonus), marks in zip(part_specs, splits):
                questions.append(AnalyzedQuestion(plabel, pcos, marks, is_bonus=pbonus))
            continue

        questions.append(AnalyzedQuestion(label, parent_cos, parent_marks, is_bonus=parent_bonus))
    return questions


def scale_questions(
    questions: list[AnalyzedQuestion],
    *,
    paper_total_marks: float,
    weightage: float,
) -> list[AnalyzedQuestion]:
    if paper_total_marks <= 0 or weightage <= 0:
        return questions
    factor = weightage / paper_total_marks
    scaled: list[AnalyzedQuestion] = []
    for q in questions:
        scaled.append(
            AnalyzedQuestion(
                q.label,
                q.co_labels,
                round(q.max_marks * factor, 2),
                is_bonus=q.is_bonus,
            )
        )
    return scaled


async def analyze_question_paper_text(text: str) -> dict:
    snippet = text[:_MAX_TEXT_CHARS]
    paper_co_map = extract_paper_co_map(snippet)
    prompt = f"""Analyze this exam question paper and return ONLY valid JSON (no markdown) with this schema:
{{
  "component_name": "Quiz|MidSem|EndSem|Assignment|Lab|etc",
  "paper_total_marks": number,
  "questions": [
    {{
      "label": "Q1",
      "co_labels": [],
      "co_evidence": "",
      "max_marks": number,
      "is_bonus": false,
      "parts": [
        {{
          "label": "Q1a",
          "co_labels": [],
          "co_evidence": "",
          "max_marks": 0,
          "is_bonus": false
        }}
      ]
    }}
  ]
}}

Rules:
- Prefer one parent object per main question with a "parts" array for sub-parts (a,b,c / i,ii / 1,2).
- Use labels exactly as printed (Q7c not Q7c1). Never invent extra numbered suffixes.
- Parent max_marks is the total for that question (excluding bonus).
- Part max_marks: use the number ONLY when the paper prints marks for that part. If the paper does not print part-wise marks, set each part's max_marks to 0 (the server will divide the parent total equally).
- NEVER copy the parent total onto every sub-part. Example: "each question is worth 10 marks" with parts (a)(b) and no part marks → parent max_marks=10, parts max_marks=0.
- When part marks ARE printed (e.g. (a) [2 Marks] (b) [3 Marks]), use those exact values on the parts and set parent max_marks to their sum.
- Leave co_labels empty in the JSON; the server extracts COs from the paper text.
- Mark bonus questions/parts with is_bonus=true.
- paper_total_marks excludes bonus marks and is the total for this component before scaling.
- Only include questions/parts that actually appear in the paper.

Question paper text:
{snippet}
"""
    raw = await _generate_analysis_text(prompt)
    data = _parse_llm_json(raw)
    questions = _flatten_llm_questions(data, snippet)
    questions = dedupe_questions(questions)
    questions = apply_paper_cos(questions, paper_co_map)

    try:
        paper_total = float(data.get("paper_total_marks") or 0)
    except (TypeError, ValueError):
        paper_total = 0.0
    if paper_total <= 0:
        paper_total = sum(q.max_marks for q in questions if not q.is_bonus)

    questions = redistribute_sibling_marks(questions, paper_total_marks=paper_total)

    missing_co_labels = [q.label for q in questions if not q.is_bonus and not q.co_labels]
    warnings = []
    if missing_co_labels:
        warnings.append(
            "No COs found for: "
            + ", ".join(missing_co_labels)
            + ". Manually edit/add CO mappings before downloading."
        )
    return {
        "component_name": str(data.get("component_name") or "Component").strip() or "Component",
        "paper_total_marks": paper_total,
        "warnings": warnings,
        "questions": [
            {
                "label": q.label,
                "co_labels": q.co_labels,
                "co_label": q.co_label,
                "max_marks": q.max_marks,
                "is_bonus": q.is_bonus,
            }
            for q in questions
        ],
    }


async def analyze_question_paper_file(filename: str, content: bytes) -> dict:
    text = extract_document_text(filename, content)
    if not text.strip():
        raise ValueError("Could not extract text from the uploaded file.")
    return await analyze_question_paper_text(text)


def generate_component_workbook(
    *,
    component_name: str,
    questions: list[dict],
    paper_total_marks: float,
    weightage: float,
) -> bytes:
    parsed = [
        AnalyzedQuestion(
            str(q.get("label") or ""),
            _normalize_co_list(q.get("co_labels")) or _normalize_co_list(q.get("co_label")),
            float(q.get("max_marks") or 0),
            is_bonus=bool(q.get("is_bonus")),
        )
        for q in questions
        if str(q.get("label") or "").strip()
    ]
    scaled = scale_questions(parsed, paper_total_marks=paper_total_marks, weightage=weightage)
    from app.copo.services.marks_template_builder import AnalyzedQuestionSpec

    specs = [
        AnalyzedQuestionSpec(q.label, q.co_label, q.max_marks, is_bonus=q.is_bonus) for q in scaled
    ]
    return build_analyzed_component_workbook(
        component_name=component_name,
        questions=specs,
    )
